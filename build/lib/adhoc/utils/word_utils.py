from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from ..db.database import get_codebase_summary
import re

def render_word_document(explanations, codebase_summary,output_filename, author_name):
    document = Document()

    # Set default font size
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Add Title
    document.add_heading('Project Documentation', 0)
    # Add Author Name
    document.add_heading(f'{author_name}', level=1)
    
    # Add Codebase Summary
    document.add_heading('Codebase Summary', level=1)
    document.add_paragraph(codebase_summary)

    # Add Change Explanations
    document.add_heading('Change Explanations', level=1)

    for item in explanations:
        document.add_heading(f"File: {item['file_path']}", level=2)

        if 'commit_message' in item:
            document.add_paragraph(f"Commit Message: {item['commit_message']}", style='List Bullet')

        document.add_paragraph('Explanation:', style='List Bullet')
        document.add_paragraph(item['explanation'])

        # Add a horizontal line
        document.add_page_break()

    # Save the document
    document.save(output_filename)
